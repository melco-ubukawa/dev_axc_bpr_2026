from flask import Blueprint, render_template, request, jsonify
from .utils.auth import get_entra_user_info
from .utils.design_project_service import DesignProjectService
from .utils.cosmos_db import get_projects_by_user, get_design_projects_by_user
from .utils.design_ai import DesignAIHandler
from .utils.diagram_gen import DiagramGenerator
import logging
import threading
import uuid
import time
import io
import os
from docx import Document
from xhtml2pdf import pisa
from flask import send_file
import markdown
import zipfile
from io import BytesIO

FONT_PATH = os.path.join(os.path.dirname(__file__), 'static', 'fonts', 'ipaexg.ttf')

logger = logging.getLogger(__name__)

# Blueprintの定義
design_bp = Blueprint('design', __name__)

# 簡易的なタスク状態管理 (本番ではRedis等を使用すべき)
auto_gen_tasks = {}

# --- ページ表示ルート ---

@design_bp.route('/design_assistant')
def design_top():
    """設計アシスタントのトップページ（ダッシュボード）を表示"""
    return render_template('dev_specs_top.html')

@design_bp.route('/design_assistant/level1/<project_id>')
def design_level1(project_id):
    """Level 1 ウィザード画面を表示"""
    return render_template('dev_specs_setup.html', project_id=project_id)

@design_bp.route('/design_assistant/editor/<project_id>')
def design_editor(project_id):
    """Level 2 詳細エディタ画面を表示"""
    return render_template('dev_specs_editor.html', project_id=project_id)


# --- APIエンドポイント ---

@design_bp.route('/api/design/projects', methods=['GET'])
def list_projects_api():
    """
    ユーザーのプロジェクト一覧を取得する。
    """
    try:
        user_info = get_entra_user_info(request)
        user_id = user_info.get("userId", "anonymous")
        
        projects = get_design_projects_by_user(user_id)
        
        return jsonify({"success": True, "projects": projects})
            
    except Exception as e:
        logger.error(f"List projects API error: {e}", exc_info=True)
        return jsonify({"error": str(e)}), 500

@design_bp.route('/api/design/create', methods=['POST'])
def create_project_api():
    """
    新規プロジェクトを作成する。
    """
    try:
        data = request.get_json()
        if not data:
            return jsonify({"error": "No JSON data provided"}), 400

        project_name = data.get('projectName')
        level = str(data.get('developerLevel', '1'))
        initial_idea = data.get('initialIdea')

        if not all([project_name, initial_idea]):
             return jsonify({"error": "Project name and idea are required"}), 400

        user_info = get_entra_user_info(request)
        user_id = user_info.get("userId", "anonymous")

        # ファイル名リストを受け取る
        file_names = data.get('initialContextFiles', [])
        context_files_data = [{"name": name, "status": "pending_upload"} for name in file_names]

        initial_project_data = {
            "documents": {},
            "diagrams": {},
            "chatHistory": [],
            "contextFiles": context_files_data,
            "initialIdea": initial_idea  # initialIdeaを保存
        }
        
        # Level 2の場合のみ、選択されたドキュメントの空エントリを作成
        selected_docs = data.get('selectedDocs', [])
        if level == '2' and selected_docs:
            for doc_id in selected_docs:
                initial_project_data["documents"][doc_id] = ""
        
        project = DesignProjectService.create_project(user_id, project_name, level, initial_project_data)
        
        if project:
            return jsonify({"success": True, "projectId": project["id"]})
        else:
            return jsonify({"error": "Failed to create project"}), 500

    except Exception as e:
        logger.error(f"Create project API error: {e}", exc_info=True)
        return jsonify({"error": str(e)}), 500


@design_bp.route('/api/design/project/<project_id>', methods=['GET'])
def get_project_api(project_id):
    """
    プロジェクトの全データを取得する。
    """
    try:
        user_info = get_entra_user_info(request)
        user_id = user_info.get("userId", "anonymous")
        
        project = DesignProjectService.get_project(project_id, user_id)
        
        if project:
            return jsonify({"success": True, "project": project})
        else:
            return jsonify({"error": "Project not found"}), 404
            
    except Exception as e:
        logger.error(f"Get project API error: {e}", exc_info=True)
        return jsonify({"error": str(e)}), 500


@design_bp.route('/api/design/chat', methods=['POST'])
def chat_api():
    """
    ユーザーからのチャットメッセージを受け取り、AI応答とドキュメント更新を行う。
    """
    try:
        data = request.get_json()
        project_id = data.get('projectId')
        user_message = data.get('message')
        doc_id = data.get('docId')
        
        if not project_id or not user_message or not doc_id:
            return jsonify({"error": "Missing projectId, message or docId"}), 400
        
        user_info = get_entra_user_info(request)
        user_id = user_info.get("userId", "anonymous")

        project = DesignProjectService.get_project(project_id, user_id)
        if not project:
            return jsonify({"error": "Project not found"}), 404

        current_doc_content = project.get("documents", {}).get(doc_id, "")
        chat_history = project.get("chatHistory", [])
        chat_history.append({"role": "user", "content": user_message})

        ai_handler = DesignAIHandler(user_info)
        
        ai_result = ai_handler.chat_and_update(
            current_doc_id=doc_id,
            current_doc_content=current_doc_content,
            chat_history=chat_history,
            user_message=user_message,
            developer_level=project.get("developerLevel", "Level 1")
        )

        new_assistant_message = ai_result.get("assistant_message", "申し訳ありません、応答を生成できませんでした。")
        updated_doc = ai_result.get("updated_document")
        diagram_code_obj = ai_result.get("diagram_code")

        chat_history.append({"role": "assistant", "content": new_assistant_message})
        
        updates = {
            "chatHistory": chat_history,
            "documents": project.get("documents", {}),
            "diagrams": project.get("diagrams", {})
        }
        
        if updated_doc is not None:
            updates["documents"][doc_id] = updated_doc
        
        if diagram_code_obj and diagram_code_obj.get("code"):
            updates["diagrams"][doc_id] = diagram_code_obj

        DesignProjectService.update_project(project_id, user_id, updates)
        
        return jsonify({
            "success": True,
            "assistant_message": new_assistant_message,
            "updated_document": updated_doc,
            "diagram_code": diagram_code_obj
        })

    except Exception as e:
        logger.error(f"Chat API error: {e}", exc_info=True)
        return jsonify({"error": str(e)}), 500

@design_bp.route('/api/design/generate_draft', methods=['POST'])
def generate_draft_api():
    """
    指定ドキュメントの初期ドラフトをAI生成する。
    インフォグラフィックが含まれる場合はそれも保存・返却する。
    """
    try:
        data = request.get_json()
        project_id = data.get('projectId')
        doc_id = data.get('docId')
        
        user_info = get_entra_user_info(request)
        user_id = user_info.get("userId", "anonymous")

        project = DesignProjectService.get_project(project_id, user_id)
        if not project: return jsonify({"error": "Project not found"}), 404

        ai_handler = DesignAIHandler(user_info)
        
        initial_idea = project.get("documents", {}).get("000_010", project.get("initialIdea", project.get("projectName")))
        
        # ▼▼▼ 修正: 全ドキュメントを渡す ▼▼▼
        result = ai_handler.generate_document_draft(
            doc_id, 
            project.get("projectName"), 
            initial_idea,
            project.get("developerLevel", "Level 2"),
            all_documents=project.get("documents", {}) # 依存関係解決のため全Docsを渡す
        )
        # ▲▲▲ 修正終了 ▲▲▲

        updates = { 
            "documents": project.get("documents", {}), 
            "diagrams": project.get("diagrams", {}),
            "infographics": project.get("infographics", {})
        }
        
        if result.get("content"):
            updates["documents"][doc_id] = result.get("content")
            
        if result.get("diagram_code") and result["diagram_code"].get("code"):
            updates["diagrams"][doc_id] = result["diagram_code"]

        if result.get("infographic"):
            updates["infographics"][doc_id] = result["infographic"]
        
        DesignProjectService.update_project(project_id, user_id, updates)
        
        return jsonify({"success": True, "result": result})

    except Exception as e:
        logger.error(f"Generate draft error: {e}", exc_info=True)
        return jsonify({"error": str(e)}), 500

@design_bp.route('/api/render/<diagram_type>', methods=['POST'])
def render_diagram_api(diagram_type):
    """
    図解コードを受け取り、SVG(またはHTML)を生成して返す。
    PlantUML, Graphviz に対応。
    """
    try:
        data = request.get_json()
        code = data.get('code')
        
        if not code:
            return jsonify({"error": "No code provided"}), 400

        user_info = get_entra_user_info(request)
        diagram_gen = DiagramGenerator(user_info)
        
        # DiagramGenerator.render_svg は SVG文字列 を返す
        svg_output = diagram_gen.render_svg(diagram_type, code)
        
        if svg_output:
            return jsonify({"success": True, "svg": svg_output})
        else:
            return jsonify({"error": "Failed to render diagram"}), 500
            
    except Exception as e:
        logger.error(f"Render diagram error: {e}", exc_info=True)
        return jsonify({"error": str(e)}), 500

@design_bp.route('/api/design/documents/definitions', methods=['GET'])
def get_document_definitions_api():
    """
    設計ドキュメントの全定義一覧を返すAPI。
    Level 2のプロジェクト作成モーダルで使用する。
    """
    try:
        from .utils.design_ai import DOCUMENT_DEFINITIONS
        # design_ai.py からインポートしたドキュメント定義をそのままJSONで返す
        return jsonify({"success": True, "definitions": DOCUMENT_DEFINITIONS})
    except Exception as e:
        logger.error(f"Get document definitions API error: {e}", exc_info=True)
        return jsonify({"error": str(e)}), 500

@design_bp.route('/api/design/auto_generate', methods=['POST'])
def auto_generate_api():
    """
    全自動生成を開始する。
    """
    try:
        data = request.get_json()
        project_id = data.get('projectId')
        doc_ids = data.get('docIds', [])
        
        user_info = get_entra_user_info(request)
        user_id = user_info.get("userId", "anonymous")

        if not project_id or not doc_ids:
            return jsonify({"error": "Missing projectId or docIds"}), 400

        task_id = f"gen_{project_id}_{uuid.uuid4().hex[:8]}"
        auto_gen_tasks[task_id] = {
            "status": "running",
            "progress": 0,
            "total": len(doc_ids),
            "current_doc": "",
            "logs": []
        }

        def background_generate(tid, pid, uid, docs):
            try:
                # 初期状態のプロジェクトを取得
                project = DesignProjectService.get_project(pid, uid)
                base_idea = project.get("documents", {}).get("000_010", project.get("initialIdea", ""))
                
                # ループ中に更新されていくドキュメントの内容を保持するローカル変数
                current_documents = project.get("documents", {}).copy()

                ai_handler = DesignAIHandler({"userId": uid})

                for i, doc_id in enumerate(docs):
                    if auto_gen_tasks[tid]["status"] == "stopped": break
                    
                    auto_gen_tasks[tid]["current_doc"] = doc_id
                    auto_gen_tasks[tid]["logs"].append(f"Generating {doc_id}...")
                    
                    # ▼▼▼ 修正: 最新のドキュメント状態を渡す ▼▼▼
                    result = ai_handler.generate_document_draft(
                        doc_id, 
                        project.get("projectName"), 
                        base_idea,
                        project.get("developerLevel", "Level 2"),
                        all_documents=current_documents # 更新されたドキュメント辞書を渡す
                    )
                    # ▲▲▲ 修正終了 ▲▲▲
                    
                    # 生成結果をローカル変数に反映（次のループで参照されるように）
                    if result.get("content"):
                        current_documents[doc_id] = result.get("content")

                    # DB更新用辞書
                    updates = { 
                        "documents": { doc_id: result.get("content", "") },
                        "diagrams": {},
                        "infographics": {} 
                    }
                    
                    if result.get("diagram_code"):
                         updates["diagrams"][doc_id] = result.get("diagram_code")

                    if result.get("infographic"):
                        updates["infographics"][doc_id] = result.get("infographic")

                    DesignProjectService.update_project(pid, uid, updates)
                    
                    auto_gen_tasks[tid]["progress"] = i + 1
                    time.sleep(2) 

                auto_gen_tasks[tid]["status"] = "completed"
                auto_gen_tasks[tid]["logs"].append("All done.")

            except Exception as e:
                logger.error(f"Auto generation failed: {e}")
                auto_gen_tasks[tid]["status"] = "error"
                auto_gen_tasks[tid]["logs"].append(str(e))

        thread = threading.Thread(target=background_generate, args=(task_id, project_id, user_id, doc_ids))
        thread.start()

        return jsonify({"success": True, "taskId": task_id})

    except Exception as e:
        logger.error(f"Auto generate API error: {e}", exc_info=True)
        return jsonify({"error": str(e)}), 500

@design_bp.route('/api/design/task_status/<task_id>', methods=['GET'])
def get_task_status_api(task_id):
    """タスクの進捗状況を返す"""
    task = auto_gen_tasks.get(task_id)
    if task:
        return jsonify(task)
    else:
        return jsonify({"error": "Task not found"}), 404

@design_bp.route('/api/design/stop_task/<task_id>', methods=['POST'])
def stop_task_api(task_id):
    """タスクを停止する"""
    if task_id in auto_gen_tasks:
        auto_gen_tasks[task_id]["status"] = "stopped"
        return jsonify({"success": True})
    return jsonify({"error": "Task not found"}), 404

@design_bp.route('/api/design/project/<project_id>', methods=['DELETE'])
def delete_project_api(project_id):
    """
    【追加】プロジェクトを削除するAPI
    """
    try:
        user_info = get_entra_user_info(request)
        user_id = user_info.get("userId", "anonymous")
        
        # サービス層の削除メソッドを呼び出し
        success = DesignProjectService.delete_project(project_id, user_id)
        
        if success:
            return jsonify({"success": True})
        else:
            return jsonify({"error": "Failed to delete project"}), 500
            
    except Exception as e:
        logger.error(f"Delete project API error: {e}", exc_info=True)
        return jsonify({"error": str(e)}), 500

@design_bp.route('/api/design/export', methods=['POST'])
def export_document_api():
    """
    ドキュメントのエクスポート処理 (Word/Markdown)
    Word出力時にMarkdownの表を検出し、枠線付きのテーブルとして出力する機能を強化。
    """
    import base64
    from docx.shared import Inches

    try:
        data = request.get_json()
        format_type = data.get('format', 'md')
        content_md = data.get('content', '')
        base_filename = data.get('filename', 'document')
        
        infographic_b64 = data.get('infographic_image')
        diagram_b64 = data.get('diagram_image')

        if not content_md:
            return jsonify({"error": "No content provided"}), 400

        file_stream = io.BytesIO()
        mimetype = 'text/plain' 
        download_name = base_filename

        if format_type == 'docx':
            doc = Document()
            doc.add_heading(base_filename, 0)

            # 1. インフォグラフィック追加
            if infographic_b64:
                try:
                    if "," in infographic_b64:
                        header, encoded = infographic_b64.split(",", 1)
                    else:
                        encoded = infographic_b64
                    img_bytes = base64.b64decode(encoded)
                    doc.add_picture(io.BytesIO(img_bytes), width=Inches(6))
                    doc.add_paragraph("") 
                except Exception as e:
                    logger.warning(f"Failed to add infographic to docx: {e}")

            # 2. 本文解析と書き込み (表対応版)
            lines = content_md.split('\n')
            i = 0
            in_code_block = False

            while i < len(lines):
                line = lines[i]
                stripped = line.strip()

                # コードブロック判定
                if stripped.startswith('```'):
                    in_code_block = not in_code_block
                    p = doc.add_paragraph(line)
                    p.style = 'No Spacing'
                    i += 1
                    continue
                
                if in_code_block:
                    p = doc.add_paragraph(line)
                    p.style = 'No Spacing'
                    i += 1
                    continue

                # --- 表(Table)の検出と処理 ---
                # パイプで始まりパイプで終わる行をテーブルとみなす
                if stripped.startswith('|') and stripped.endswith('|'):
                    table_lines = []
                    # 連続するテーブル行を取得
                    while i < len(lines) and lines[i].strip().startswith('|'):
                        table_lines.append(lines[i].strip())
                        i += 1
                    
                    # テーブル処理 (最低2行：ヘッダーとセパレータが必要)
                    if len(table_lines) >= 2:
                        try:
                            # 1行目をヘッダーとして解析
                            headers = [c.strip() for c in table_lines[0].strip('|').split('|')]
                            
                            # 2行目はセパレータ (|---|) なのでスキップ
                            
                            # 3行目以降はデータ
                            rows_data = []
                            for r_line in table_lines[2:]:
                                # セルの分割
                                cells = [c.strip() for c in r_line.strip('|').split('|')]
                                rows_data.append(cells)

                            # Wordのテーブルを作成
                            table = doc.add_table(rows=1, cols=len(headers))
                            
                            # ★★★ ここが重要: 枠線付きのスタイルを適用 ★★★
                            table.style = 'Table Grid' 

                            # ヘッダー書き込み
                            hdr_cells = table.rows[0].cells
                            for idx, text in enumerate(headers):
                                if idx < len(hdr_cells):
                                    hdr_cells[idx].text = text
                            
                            # データ行書き込み
                            for r_data in rows_data:
                                row_cells = table.add_row().cells
                                for idx, text in enumerate(r_data):
                                    if idx < len(row_cells):
                                        row_cells[idx].text = text
                        except Exception as tbl_e:
                            logger.warning(f"Failed to parse markdown table: {tbl_e}")
                            # パース失敗時はテキストとして出力
                            for tl in table_lines:
                                doc.add_paragraph(tl, style='No Spacing')
                    else:
                        # 行数が足りない場合はそのままテキスト出力
                        for tl in table_lines:
                            doc.add_paragraph(tl)
                    
                    # ループ処理ですでに i は進んでいるので continue
                    continue

                # --- 通常のMarkdown要素処理 ---
                if line.startswith('# '):
                    doc.add_heading(line[2:], level=1)
                elif line.startswith('## '):
                    doc.add_heading(line[3:], level=2)
                elif line.startswith('### '):
                    doc.add_heading(line[4:], level=3)
                elif line.startswith('- ') or line.startswith('* '):
                    doc.add_paragraph(line[2:], style='List Bullet')
                elif line.startswith('1. '):
                    doc.add_paragraph(line[3:], style='List Number')
                else:
                    doc.add_paragraph(line)
                
                i += 1

            # 3. 図解追加
            if diagram_b64:
                try:
                    doc.add_page_break()
                    doc.add_heading('Diagram / Visualization', level=2)
                    if "," in diagram_b64:
                        header, encoded = diagram_b64.split(",", 1)
                    else:
                        encoded = diagram_b64
                    img_bytes = base64.b64decode(encoded)
                    doc.add_picture(io.BytesIO(img_bytes), width=Inches(6.5))
                except Exception as e:
                    logger.warning(f"Failed to add diagram to docx: {e}")

            doc.save(file_stream)
            mimetype = 'application/vnd.openxmlformats-officedocument.wordprocessingml.document'
            download_name += '.docx'

        elif format_type == 'md':
            file_stream.write(content_md.encode('utf-8'))
            mimetype = 'text/markdown'
            download_name += '.md'

        else:
            return jsonify({"error": f"Unsupported format: {format_type}"}), 400

        file_stream.seek(0)
        return send_file(
            file_stream,
            as_attachment=True,
            download_name=download_name,
            mimetype=mimetype
        )

    except Exception as e:
        logger.error(f"Export error: {e}", exc_info=True)
        return jsonify({"error": str(e)}), 500

@design_bp.route('/api/design/export_all/<project_id>', methods=['GET'])
def export_all_documents_api(project_id):
    """
    (2) プロジェクト内の全ドキュメントを一括でZIPダウンロードするAPI
    """
    import zipfile
    from io import BytesIO
    from docx import Document
    from docx.shared import Inches
    
    from .utils.design_ai import DOCUMENT_DEFINITIONS
    from .utils.diagram_gen import DiagramGenerator
    from .utils.storage import download_bytes_from_blob_url # ★追加

    try:
        user_info = get_entra_user_info(request)
        user_id = user_info.get("userId", "anonymous")
        
        project = DesignProjectService.get_project(project_id, user_id)
        if not project:
            return jsonify({"error": "Project not found"}), 404
            
        documents = project.get("documents", {})
        diagrams = project.get("diagrams", {})
        infographics = project.get("infographics", {})
        project_name = project.get("projectName", "project")
        
        diagram_gen = DiagramGenerator(user_info)

        zip_buffer = BytesIO()
        with zipfile.ZipFile(zip_buffer, 'w', zipfile.ZIP_DEFLATED) as zip_file:
            
            for doc_id, content in documents.items():
                if not content: continue
                
                doc_def = DOCUMENT_DEFINITIONS.get(doc_id, {})
                doc_title = doc_def.get("title", doc_id)
                safe_title = "".join([c for c in doc_title if c.isalnum() or c in (' ', '-', '_')]).strip()
                base_filename = f"{doc_id}_{safe_title}"
                
                # ==========================================
                # 0. アセット準備 (インフォグラフィック)
                # ==========================================
                infographic_bytes = None
                if doc_id in infographics:
                    img_url = infographics[doc_id]
                    if img_url:
                        try:
                            # ★修正: HTTPリクエストではなく、内部ストレージロジックで取得
                            img_data = download_bytes_from_blob_url(img_url)
                            if img_data:
                                infographic_bytes = BytesIO(img_data)
                                zip_file.writestr(f"images/{base_filename}_infographic.png", img_data)
                        except Exception as img_err:
                            logger.warning(f"Failed to fetch infographic for {doc_id}: {img_err}")

                # ... (以下、図解・Markdown・Word出力ロジックは変更なし) ...
                
                # ==========================================
                # 1. 図解ファイルの保存 (diagramsフォルダ)
                # ==========================================
                if doc_id in diagrams:
                    diag = diagrams[doc_id]
                    if diag and diag.get("code"):
                        d_type = diag.get("type", "txt")
                        d_code = diag.get("code")
                        
                        ext_map = {
                            "plantuml": "puml", "graphviz": "dot", "markmap": "md", 
                            "html": "html", "svg": "svg", "mermaid": "mmd"
                        }
                        src_ext = ext_map.get(d_type, "txt")
                        
                        # ソースコードの保存
                        zip_file.writestr(f"diagrams/{base_filename}.{src_ext}", d_code)

                        # サーバー側レンダリング (PlantUML, Graphviz) -> SVG保存
                        if d_type in ["plantuml", "graphviz"]:
                            try:
                                svg_out = diagram_gen.render_svg(d_type, d_code)
                                if svg_out:
                                    zip_file.writestr(f"diagrams/{base_filename}.svg", svg_out)
                            except Exception as render_err:
                                logger.warning(f"Failed to render diagram for export {doc_id}: {render_err}")
                        
                        # Markmap -> Standalone HTML
                        if d_type == "markmap":
                            markmap_html = f"""
<!DOCTYPE html>
<html>
<head>
<meta charset="UTF-8">
<title>{doc_title} - Markmap</title>
<style>svg.markmap {{ width: 100%; height: 100vh; }}</style>
<script src="https://cdn.jsdelivr.net/npm/d3@7"></script>
<script src="https://cdn.jsdelivr.net/npm/markmap-view@0.15.4"></script>
<script src="https://cdn.jsdelivr.net/npm/markmap-lib@0.15.4/dist/browser/index.min.js"></script>
</head>
<body>
<svg id="mindmap" class="markmap"></svg>
<script>
(async () => {{
    const markdown = `{d_code.replace('`', '\\`').replace('$', '\\$')}`;
    const {{ Transformer, Markmap, loadCSS, loadJS }} = window.markmap;
    const transformer = new Transformer();
    const {{ root, features }} = transformer.transform(markdown);
    const {{ styles, scripts }} = transformer.getUsedAssets(features);
    if (styles) loadCSS(styles);
    if (scripts) loadJS(scripts, {{ getMarkmap: () => Markmap }});
    Markmap.create('#mindmap', null, root);
}})();
</script>
</body>
</html>
"""
                            zip_file.writestr(f"diagrams/{base_filename}_markmap.html", markmap_html)

                        if d_type == "html" and src_ext != "html":
                             zip_file.writestr(f"diagrams/{base_filename}.html", d_code)


                # ==========================================
                # 2. Markdownファイルの作成
                # ==========================================
                md_content = content
                if doc_id in diagrams:
                    diag = diagrams[doc_id]
                    if diag and diag.get("code"):
                         md_content += f"\n\n## Diagram Code ({diag.get('type')})\n\n```{diag.get('type')}\n{diag.get('code')}\n```"
                
                zip_file.writestr(f"{base_filename}.md", md_content)

                # ==========================================
                # 3. Wordファイルの作成
                # ==========================================
                try:
                    doc = Document()
                    doc.add_heading(doc_title, 0)

                    # インフォグラフィック画像の埋め込み
                    if infographic_bytes:
                        try:
                            infographic_bytes.seek(0)
                            doc.add_picture(infographic_bytes, width=Inches(6))
                            doc.add_paragraph("Project Concept Infographic", style="Caption")
                        except Exception as e:
                            logger.warning(f"Failed to embed infographic to docx: {e}")

                    # 本文のパースと書き込み
                    in_code_block = False
                    for line in content.split('\n'):
                        stripped = line.strip()
                        if stripped.startswith('```'):
                            in_code_block = not in_code_block
                            p = doc.add_paragraph(line)
                            p.style = 'No Spacing'
                            continue
                        if in_code_block:
                            p = doc.add_paragraph(line)
                            p.style = 'No Spacing'
                            continue
                        
                        if line.startswith('# '): doc.add_heading(line[2:], level=1)
                        elif line.startswith('## '): doc.add_heading(line[3:], level=2)
                        elif line.startswith('### '): doc.add_heading(line[4:], level=3)
                        elif line.startswith('- ') or line.startswith('* '): doc.add_paragraph(line[2:], style='List Bullet')
                        elif line.startswith('1. '): doc.add_paragraph(line[3:], style='List Number')
                        else: doc.add_paragraph(line)

                    # 図解コードの書き込み
                    if doc_id in diagrams:
                        diag = diagrams[doc_id]
                        if diag and diag.get("code"):
                            doc.add_page_break()
                            doc.add_heading(f"Diagram Source ({diag.get('type')})", level=2)
                            doc.add_paragraph(f"※図解ファイルは diagrams/{base_filename}.* を参照してください。", style="Quote")
                            code_p = doc.add_paragraph(diag.get('code'))
                            code_p.style = 'No Spacing'

                    docx_buffer = BytesIO()
                    doc.save(docx_buffer)
                    zip_file.writestr(f"{base_filename}.docx", docx_buffer.getvalue())

                except Exception as docx_err:
                    logger.error(f"Failed to generate docx for {doc_id}: {docx_err}")

        zip_buffer.seek(0)
        
        return send_file(
            zip_buffer,
            as_attachment=True,
            download_name=f"{project_name}_documents.zip",
            mimetype='application/zip'
        )

    except Exception as e:
        logger.error(f"Export all API error: {e}", exc_info=True)
        return jsonify({"error": str(e)}), 500
