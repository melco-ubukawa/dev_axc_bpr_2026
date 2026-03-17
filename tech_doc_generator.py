"""
技術文書自動生成アプリケーション（修正版）
AWS_BEARER_TOKEN_BEDROCK を正しく使用
"""

import streamlit as st
from docx import Document
import anthropic
import io
from datetime import datetime
import PyPDF2
from pathlib import Path
import os

# ページ設定
st.set_page_config(
    page_title="技術文書自動生成ツール",
    page_icon="📝",
    layout="wide",
    initial_sidebar_state="expanded"
)

# カスタムCSS
st.markdown("""
<style>
    .main-header {
        font-size: 2.5rem;
        color: #1f77b4;
        text-align: center;
        padding: 1rem;
        background: linear-gradient(90deg, #e3f2fd 0%, #bbdefb 100%);
        border-radius: 10px;
        margin-bottom: 2rem;
    }
    .step-header {
        font-size: 1.5rem;
        color: #2e7d32;
        padding: 0.5rem;
        background-color: #f1f8e9;
        border-left: 5px solid #4caf50;
        margin: 1rem 0;
    }
    .info-box {
        background-color: #e3f2fd;
        padding: 1rem;
        border-radius: 5px;
        margin: 1rem 0;
    }
</style>
""", unsafe_allow_html=True)


def extract_text_from_pdf(pdf_file):
    """PDFファイルからテキストを抽出"""
    try:
        pdf_reader = PyPDF2.PdfReader(pdf_file)
        text = ""
        for page in pdf_reader.pages:
            text += page.extract_text() + "\n"
        return text
    except Exception as e:
        st.error(f"PDFの読み込みエラー: {e}")
        return None


def extract_text_from_docx(docx_file):
    """Wordファイルからテキストを抽出"""
    try:
        doc = Document(docx_file)
        text = "\n".join([para.text for para in doc.paragraphs if para.text.strip()])
        return text
    except Exception as e:
        st.error(f"Wordファイルの読み込みエラー: {e}")
        return None


def create_bedrock_client():
    """Bedrock クライアントを作成"""
    import boto3
    aws_region = 'us-west-2'
    return boto3.client('bedrock-runtime', region_name=aws_region)


def create_anthropic_client(api_key):
    """Anthropic クライアントを作成"""
    return anthropic.Anthropic(api_key=api_key, timeout=60.0)


def _call_bedrock_api(bedrock_client, prompt):
    """Bedrock API を呼び出し"""
    import json

    try:
        model_id = 'us.anthropic.claude-opus-4-6-v1'

        response = bedrock_client.invoke_model(
            modelId=model_id,
            body=json.dumps({
                "anthropic_version": "bedrock-2023-05-31",
                "max_tokens": 2000,
                "messages": [{"role": "user", "content": prompt}]
            }),
            contentType='application/json',
            accept='application/json'
        )

        response_body = json.loads(response['body'].read())
        content = response_body.get('content', [])

        if not content:
            st.error("❌ レスポンスが空です")
            with st.expander("🔍 詳細"):
                st.json(response_body)
            return None

        return content[0].get('text', '')

    except Exception as e:
        st.error(f"❌ API エラー: {str(e)}")
        with st.expander("🔍 詳細"):
            import traceback
            st.code(traceback.format_exc())
        return None


def _call_anthropic_api(client, prompt):
    """Anthropic API を呼び出し"""
    try:
        model_id = "claude-sonnet-4-20250514"

        message = client.messages.create(
            model=model_id,
            max_tokens=2000,
            messages=[{"role": "user", "content": prompt}]
        )

        if not message or not hasattr(message, 'content') or not message.content:
            st.error("❌ レスポンスが空です")
            return None

        first_block = message.content[0]

        if hasattr(first_block, 'text'):
            return first_block.text
        elif isinstance(first_block, dict) and 'text' in first_block:
            return first_block['text']
        else:
            st.error(f"❌ 想定外のレスポンス形式")
            return None

    except Exception as e:
        st.error(f"❌ API エラー: {str(e)}")
        with st.expander("🔍 詳細"):
            import traceback
            st.code(traceback.format_exc())
        return None


def generate_summary_and_outline(reference_text, api_key):
    """Claude APIで大要とアウトラインを生成"""
    try:
        # 設定
        api_provider = 'bedrock'
        aws_region = 'us-west-2'
        bearer_token = os.environ.get('AWS_BEARER_TOKEN_BEDROCK')

        # クライアント作成
        if api_provider == 'bedrock':
            if not bearer_token:
                st.error("❌ AWS_BEARER_TOKEN_BEDROCK が設定されていません")
                return None, None

            try:
                import json
                bedrock_client = create_bedrock_client()
            except Exception as e:
                st.error(f"❌ 接続エラー: {str(e)}")
                with st.expander("🔍 詳細情報"):
                    import traceback
                    st.code(traceback.format_exc())
                return None, None
        else:
            if not api_key:
                st.error("❌ API Keyが設定されていません")
                return None, None
            client = create_anthropic_client(api_key)

        # 参考資料を切り詰める
        max_text_length = 10000
        truncated_text = reference_text[:max_text_length]
        if len(reference_text) > max_text_length:
            st.warning(f"⚠️ 参考資料が長いため、先頭{max_text_length}文字のみを使用します")

        # プロンプト作成
        prompt = f"""以下の参考資料を元に、技術文書の「大要」と「目次（アウトライン）」を生成してください。

【参考資料】
{truncated_text}

【出力形式】
以下の形式で厳密に出力してください：

## 大要
（200-300文字程度で、この文書の目的、対象、主要な内容を簡潔にまとめる）

## 目次
1. はじめに
2. （適切な章立て）
3. （適切な章立て）
4. まとめ

技術文書として適切な構成と表現を心がけてください。
必ず「## 大要」と「## 目次」の見出しを含めてください。"""

        # API リクエスト
        if api_provider == 'bedrock':
            result = _call_bedrock_api(bedrock_client, prompt)
        else:
            result = _call_anthropic_api(client, prompt)

        if not result:
            return None, None

        # 大要と目次を分離
        if "## 目次" in result:
            parts = result.split("## 目次")
            summary = parts[0].replace("## 大要", "").strip()
            outline = "## 目次\n" + parts[1].strip() if len(parts) > 1 else ""
        elif "目次" in result and "大要" in result:
            # 異なる形式の場合
            lines = result.split('\n')
            summary_lines = []
            outline_lines = []
            in_summary = False
            in_outline = False

            for line in lines:
                if '大要' in line:
                    in_summary = True
                    in_outline = False
                    continue
                elif '目次' in line:
                    in_summary = False
                    in_outline = True
                    outline_lines.append(line)
                    continue

                if in_summary:
                    if line is not None:
                        summary_lines.append(line)
                elif in_outline:
                    if line is not None:
                        outline_lines.append(line)

            summary_lines_clean = [str(l) for l in summary_lines if l is not None]
            outline_lines_clean = [str(l) for l in outline_lines if l is not None]

            summary = '\n'.join(summary_lines_clean).strip()
            outline = '\n'.join(outline_lines_clean).strip()
        else:
            # フォーマットが異なる場合
            st.warning("⚠️ 想定外のフォーマットです。手動で分割してください。")
            with st.expander("📄 生成されたテキスト全体"):
                st.text(result)
            summary = result[:500] if len(result) > 500 else result
            outline = ""

        return summary, outline

    except anthropic.APIConnectionError as e:
        st.error(f"❌ API接続エラー: ネットワーク接続を確認してください")
        st.error(f"詳細: {str(e)}")
        return None, None
    except anthropic.AuthenticationError as e:
        st.error(f"❌ 認証エラー: API Keyまたはトークンが正しいか確認してください")
        st.error(f"詳細: {str(e)}")
        return None, None
    except anthropic.PermissionDeniedError as e:
        st.error(f"❌ 権限エラー: トークンに必要な権限がありません")
        st.error(f"詳細: {str(e)}")
        st.info("💡 AWS_BEARER_TOKEN_BEDROCK が正しいか確認してください")
        return None, None
    except Exception as e:
        st.error(f"❌ AI生成エラー: {type(e).__name__}")
        st.error(f"詳細: {str(e)}")
        with st.expander("🔍 詳細なエラー情報"):
            import traceback
            st.code(traceback.format_exc())
        return None, None


def fill_word_template(template_path, data):
    """Wordテンプレートにデータを埋め込む"""
    try:
        doc = Document(template_path)

        # テーブル3: 表題、作成日
        if len(doc.tables) >= 3:
            table3 = doc.tables[2]
            if data.get("title"):
                table3.rows[0].cells[1].text = data["title"]
                table3.rows[0].cells[2].text = data["title"]

            if data.get("date"):
                date_str = data["date"].strftime("年 月 日：　%Y年　%m月%d日")
                table3.rows[1].cells[1].text = date_str

        # テーブル4: 著者情報
        if len(doc.tables) >= 4:
            table4 = doc.tables[3]
            if data.get("department"):
                table4.rows[1].cells[0].text = data["department"]
            if data.get("author"):
                table4.rows[1].cells[1].text = data["author"]
            if data.get("furigana"):
                table4.rows[1].cells[2].text = data["furigana"]
            if data.get("man_number"):
                table4.rows[1].cells[3].text = data["man_number"]

        # テーブル5: 大要と目次
        if len(doc.tables) >= 5:
            table5 = doc.tables[4]
            if data.get("summary"):
                table5.rows[1].cells[0].text = data["summary"]

            if data.get("outline"):
                table5.rows[3].cells[0].text = data["outline"]

        # テーブル6: キーワード
        if len(doc.tables) >= 6:
            table6 = doc.tables[5]
            if data.get("keyword_descriptor"):
                table6.rows[0].cells[1].text = data["keyword_descriptor"]
            if data.get("keyword_free"):
                table6.rows[1].cells[1].text = data["keyword_free"]

        return doc

    except Exception as e:
        st.error(f"テンプレート処理エラー: {e}")
        return None


def main():
    """メインアプリケーション"""

    # ヘッダー
    st.markdown('<div class="main-header">📝 技術文書自動生成ツール</div>', unsafe_allow_html=True)

    st.markdown("""
    <div class="info-box">
    <strong>✨ このツールの機能</strong><br>
    参考資料を読み込んで、AIが自動的に「大要（要約）」と「目次（アウトライン）」を生成します。<br>
    その他の情報を入力すると、技術文書の表紙が完成します。
    </div>
    """, unsafe_allow_html=True)

    # サイドバー: 設定
    with st.sidebar:
        st.header("⚙️ 設定")

        # API Provider確認
        api_provider = os.environ.get('CLAUDE_API_PROVIDER', 'anthropic').lower()
        aws_region = os.environ.get('AWS_REGION', 'us-west-2')
        bearer_token = os.environ.get('AWS_BEARER_TOKEN_BEDROCK')

        if api_provider == 'bedrock':
            st.info(f"🔧 **接続方式**: AWS Bedrock")
            st.info(f"🌏 **リージョン**: {aws_region}")
            st.markdown("---")

            if bearer_token:
                st.success("✅ **認証**: AWS_BEARER_TOKEN_BEDROCK 設定済み")
                st.code(f"トークン: {bearer_token[:20]}...({len(bearer_token)}文字)")
            else:
                st.error("❌ **認証**: AWS_BEARER_TOKEN_BEDROCK が未設定")
                st.markdown("""
                **設定方法:**
                1. システム環境変数に設定:
                   - 変数名: `AWS_BEARER_TOKEN_BEDROCK`
                   - 値: 社内システムから取得したトークン
                2. コマンドプロンプトを再起動
                3. アプリを再起動
                """)

            api_key = None  # Bedrockの場合、Anthropic API Keyは不要
        else:
            st.info(f"🔧 **接続方式**: Anthropic API (直接)")
            api_key = st.text_input("Claude API Key", type="password", help="AnthropicのAPIキーを入力してください")

        st.markdown("---")
        st.markdown("### 📚 使い方")
        st.markdown("""
        1. 参考資料をアップロードまたは入力
        2. AI生成ボタンをクリック
        3. 他の項目を入力
        4. Wordファイルを生成
        """)

    # ステップ1: 参考資料の入力
    st.markdown('<div class="step-header">📄 ステップ1: 参考資料の入力</div>', unsafe_allow_html=True)

    col1, col2 = st.columns([1, 1])

    with col1:
        st.subheader("ファイルアップロード")
        uploaded_file = st.file_uploader(
            "Word/PDFファイルをアップロード",
            type=["docx", "pdf"],
            help="参考資料となるファイルを選択してください"
        )

        reference_text = ""
        if uploaded_file:
            with st.spinner("ファイルを読み込んでいます..."):
                if uploaded_file.name.endswith(".pdf"):
                    reference_text = extract_text_from_pdf(uploaded_file)
                elif uploaded_file.name.endswith(".docx"):
                    reference_text = extract_text_from_docx(uploaded_file)

                if reference_text:
                    st.success(f"✅ ファイルを読み込みました（{len(reference_text)}文字）")

    with col2:
        st.subheader("テキスト直接入力")
        text_input = st.text_area(
            "参考資料をここに貼り付け",
            height=200,
            placeholder="参考資料のテキストをここに貼り付けてください..."
        )

        if text_input:
            reference_text = text_input

    # 参考資料プレビュー
    if reference_text:
        with st.expander("📖 参考資料プレビュー"):
            st.text(reference_text[:1000] + ("..." if len(reference_text) > 1000 else ""))

    st.markdown("---")

    # ステップ2: AI生成
    st.markdown('<div class="step-header">🤖 ステップ2: AI自動生成</div>', unsafe_allow_html=True)

    col1, col2 = st.columns([1, 3])

    with col1:
        generate_button = st.button("🚀 大要と目次を生成", type="primary", use_container_width=True)

    # セッションステートの初期化
    if "summary" not in st.session_state:
        st.session_state.summary = ""
    if "outline" not in st.session_state:
        st.session_state.outline = ""

    if generate_button:
        if api_provider != 'bedrock' and not api_key:
            st.error("⚠️ Claude API Keyを入力してください")
        elif not reference_text:
            st.error("⚠️ 参考資料を入力してください")
        else:
            with st.spinner("AIが大要と目次を生成しています... ⏳"):
                summary, outline = generate_summary_and_outline(reference_text, api_key)
                if summary and outline:
                    st.session_state.summary = summary
                    st.session_state.outline = outline
                    st.success("✅ 生成完了！")
                    st.balloons()

    # 生成結果の表示と編集
    col1, col2 = st.columns(2)

    with col1:
        st.subheader("📋 大要（要約）")
        st.session_state.summary = st.text_area(
            "大要",
            value=st.session_state.summary,
            height=200,
            label_visibility="collapsed",
            help="AI生成後も編集可能です"
        )

    with col2:
        st.subheader("📑 目次（アウトライン）")
        st.session_state.outline = st.text_area(
            "目次",
            value=st.session_state.outline,
            height=200,
            label_visibility="collapsed",
            help="AI生成後も編集可能です"
        )

    st.markdown("---")

    # ステップ3: その他の情報入力
    st.markdown('<div class="step-header">✏️ ステップ3: その他の情報入力</div>', unsafe_allow_html=True)

    col1, col2 = st.columns(2)

    with col1:
        title = st.text_input("📌 文書タイトル", placeholder="例: システム設計書")
        department = st.text_input("🏢 事業所名／所属部課名", placeholder="例: 技術開発部")
        author = st.text_input("👤 著者名", placeholder="例: 山田 太郎")

    with col2:
        date = st.date_input("📅 作成日", value=datetime.now())
        furigana = st.text_input("🔤 フリガナ", placeholder="例: ヤマダ タロウ")
        man_number = st.text_input("🔢 マンナンバー", placeholder="例: 12345")

    col1, col2 = st.columns(2)

    with col1:
        keyword_descriptor = st.text_input("🏷️ キーワード (ディスクリプタ)", placeholder="例: DISPC")

    with col2:
        keyword_free = st.text_input("🏷️ キーワード (フリーターム)", placeholder="例: システム開発")

    st.markdown("---")

    # ステップ4: Word生成
    st.markdown('<div class="step-header">💾 ステップ4: Word文書生成</div>', unsafe_allow_html=True)

    col1, col2, col3 = st.columns([1, 1, 2])

    with col1:
        generate_word = st.button("📄 Word文書を生成", type="primary", use_container_width=True)

    if generate_word:
        if not st.session_state.summary or not st.session_state.outline:
            st.error("⚠️ 大要と目次を生成してください")
        elif not title:
            st.error("⚠️ 文書タイトルを入力してください")
        else:
            with st.spinner("Word文書を生成しています... 📝"):
                template_path = Path("技ノート_表紙.docx")

                if not template_path.exists():
                    st.error(f"⚠️ テンプレートファイルが見つかりません: {template_path}")
                else:
                    data = {
                        "title": title,
                        "date": date,
                        "department": department,
                        "author": author,
                        "furigana": furigana,
                        "man_number": man_number,
                        "summary": st.session_state.summary,
                        "outline": st.session_state.outline,
                        "keyword_descriptor": keyword_descriptor,
                        "keyword_free": keyword_free
                    }

                    doc = fill_word_template(str(template_path), data)

                    if doc:
                        # メモリ上にWordファイルを保存
                        doc_io = io.BytesIO()
                        doc.save(doc_io)
                        doc_io.seek(0)

                        st.success("✅ Word文書の生成に成功しました！")

                        # ダウンロードボタン
                        st.download_button(
                            label="⬇️ Word文書をダウンロード",
                            data=doc_io,
                            file_name=f"技術文書_{title}_{date.strftime('%Y%m%d')}.docx",
                            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                            use_container_width=True
                        )


if __name__ == "__main__":
    main()
